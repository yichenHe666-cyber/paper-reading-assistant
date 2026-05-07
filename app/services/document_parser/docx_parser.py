import logging
from pathlib import Path
from app.services.document_parser.models import ParsedDocument

logger = logging.getLogger("paper_reader")

_HEADING_LEVELS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "Title": 1,
}


def parse_docx(file_path: str) -> ParsedDocument:
    doc = ParsedDocument(raw_path=file_path)

    if file_path.lower().endswith(".doc") and not file_path.lower().endswith(".docx"):
        doc.parse_status = "failed"
        doc.parse_errors.append("不支持旧版 .doc 格式，请转换为 .docx 后重试")
        return doc

    try:
        from docx import Document as DocxDocument
    except ImportError:
        doc.parse_status = "failed"
        doc.parse_errors.append("python-docx 未安装，无法解析 DOCX")
        return doc

    try:
        docx_doc = DocxDocument(file_path)
    except Exception as e:
        doc.parse_status = "failed"
        doc.parse_errors.append(f"无法打开 DOCX 文件: {e}")
        return doc

    try:
        props = docx_doc.core_properties
        doc.title = props.title or Path(file_path).stem
        if props.author:
            doc.authors = [a.strip() for a in props.author.split(";") if a.strip()]

        doc.metadata = {
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "category": props.category or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
        }

        all_text_parts = []
        current_section = None
        current_content_parts = []

        for para in docx_doc.paragraphs:
            style_name = para.style.name if para.style else ""
            level = _HEADING_LEVELS.get(style_name, 0)
            text = para.text.strip()

            if not text:
                continue

            if level > 0:
                if current_section is not None or current_content_parts:
                    section_text = "\n\n".join(current_content_parts)
                    if current_section is None:
                        doc.sections.append(
                            {
                                "level": 0,
                                "title": "",
                                "content": section_text,
                                "type": "text",
                            }
                        )
                    else:
                        doc.sections.append(
                            {
                                "level": current_section[0],
                                "title": current_section[1],
                                "content": section_text,
                                "type": "text",
                            }
                        )

                current_section = (level, text)
                current_content_parts = []
            else:
                current_content_parts.append(text)

            all_text_parts.append(text)

        if current_section is not None:
            section_text = "\n\n".join(current_content_parts)
            doc.sections.append(
                {
                    "level": current_section[0],
                    "title": current_section[1],
                    "content": section_text,
                    "type": "text",
                }
            )
        elif current_content_parts:
            section_text = "\n\n".join(current_content_parts)
            doc.sections.append(
                {
                    "level": 0,
                    "title": "",
                    "content": section_text,
                    "type": "text",
                }
            )

        doc.content = "\n\n".join(all_text_parts)

        if not all_text_parts:
            doc.parse_status = "partial"
            doc.parse_errors.append("DOCX 文本提取为空")
        else:
            doc.parse_status = "success"

    except Exception as e:
        doc.parse_status = "failed"
        doc.parse_errors.append(f"DOCX 解析异常: {e}")

    return doc
